import docx
import numpy as np
import os

import win32print
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
import tkinter as tk
from Data_manager.src.Order_detail_manager import Order_detail_manager
from Data_manager.src.orders_manager import Order_manager
import win32api


class bill_manager:
    def __init__(self):
        self.order_detail = Order_detail_manager()
        self.order_control = Order_manager()

        self.origin_path = "Data_manager/bill_form/Form_bill.docx"
        self.parent_dir = "C:/Users/LENOVO/Desktop/bill_data/Bill_không_nợ/"
        self.deb_parent_dir = "C:/Users/LENOVO/Desktop/bill_data/Bill_nợ/"


    def ratio_value(self, value):
        box_size = 9
        ratio = abs(len(value) - box_size)
        return ratio * " "

    def fix_text(self, para):
        max_size = 46
        ratio = max_size - len(para)
        ratio_final = int(np.floor(ratio / 2))
        return ratio_final * ' '

    def get_bill(self, file_bill):
        doc = docx.api.Document(file_bill)
        all_paras = doc.paragraphs
        table = doc.tables[0]
        bill_infor = []
        for para in range(0, 5):
            text = "%s" % all_paras[para].text
            bill_infor.append("\n%s\n" % text)

        for row in table.rows:
            table = "  %s {} %s {} %s {} %s {} " % (
                row.cells[0].text, row.cells[1].text, row.cells[2].text, row.cells[3].text)
            final = table.format(self.ratio_value(row.cells[0].text), self.ratio_value(row.cells[1].text),
                                 self.ratio_value(row.cells[2].text), self.ratio_value(row.cells[3].text))
            bill_infor.append("\n%s\n" % final)

        for para in range(len(all_paras) - 3, len(all_paras)):
            text = "\n%s" % all_paras[para].text
            bill_infor.append(text)

        return bill_infor

    def add_infor(self, file_name, save_direction, file_option = None,font_size=11, text=None, option=None):

        """
        Add info to the form bill

        :param file_name:
        :param save_direction:
        :param font_size:
        :param text:
        :param file_option: default to use the default form bill and None for any text
        :param option: None to any option and end to add the ending paragraph
        :return:
        """

        if file_option == None:

            file_name = "../Data_manager/bill_form/Form_bill.docx"
            save_direction = "C:/Users/LENOVO/Desktop/bill_data/"+save_direction
            print(save_direction)
        doc = docx.Document(file_name)

        parag = doc.add_paragraph("")

        try:
            font_styles = doc.styles
            font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
            font_object = font_charstyle.font
            font_object.size = Pt(font_size)
            font_object.name = 'Calibri'

            if option == 'end':
                font_object.size = Pt(9)
                font_object.name = 'Calibri'
                parag.add_run(
                    "   Karaoke Hồng Xuyến chân thành cảm ơn quý khách!\n\t     Chào quý khách và hẹn gặp lại!",
                    style='CommentsStyle')
                parag.insert_paragraph_before("*********************************")
            else:
                parag.add_run(text, style='CommentsStyle')

        except:
            if option == 'end':
                parag.add_run(
                    "   Karaoke Hồng Xuyến chân thành cảm ơn quý khách!\n\t     Chào quý khách và hẹn gặp lại!",
                    style='CommentsStyle')
                parag.insert_paragraph_before("*********************************")
            else:
                parag.add_run(text)



        doc.save(save_direction)

    def insert_table(self, file_name, save_direction, nrows):
        document = docx.Document(file_name)
        table = document.add_table(rows=nrows, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        document.save(save_direction)

    def insert_row_and_value(self, file_name, save_direction, info):
        document = docx.Document(file_name)
        # print(document.tables)
        table = document.tables[0]

        # print(info)
        for item in info:

            # print(item)
            cells = table.add_row().cells
            # add value cell 1 and change the font size
            cells[0].text = str(item[1])
            paragraphs = cells[0].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(10)
            font.name  = 'Calibri'

            # add value cell 1 and change the font size
            cells[1].text = str(item[2])
            paragraphs = cells[1].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(10)
            font.name  = 'Calibri'

            # add value cell 1 and change the font size
            cells[2].text = str(item[3])
            paragraphs = cells[2].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(10)
            font.name  = 'Calibri'


            cells[3].text = str(item[4])
            paragraphs = cells[3].paragraphs
            paragraph = paragraphs[0]
            run_obj = paragraph.runs
            run = run_obj[0]
            font = run.font
            font.size = Pt(10)
            font.name  = 'Calibri'


        document.save(save_direction)

    def show_bill_to_label(self, widow, file):
        text = self.get_bill(file)
        for para in text:
            widow.tag_configure("center", justify='center')
            widow.tag_add("center", "1.0", "end")
            widow.insert(tk.END, para)

    def in_Bill(self, name):
        # print(name)
        win32api.ShellExecute(0, "print", name, '/d:"%s"' % win32print.GetDefaultPrinter (),".",0)

    def create_dir(self, directory):

        parent_dir = "C:/Users/LENOVO/Desktop/bill_data"
        if not self.check_dir(parent_dir):
            os.mkdir(parent_dir)
            print("Directory '% s' created" % directory)

        path = os.path.join(parent_dir, directory)
        if not self.check_dir(path):
            os.mkdir(path)
            print("Directory '% s' created" % path)

    def check_dir(self, path):
        isdir = os.path.isdir(path)
        return isdir

    def complete_bill(self, window, order_id, bill_name, date, origin_time_price,Total, Time, time_price):
        """

        :param window:
        :param order_id:
        :param bill_name:
        :param date:
        :param origin_time_price:
        :param Total:
        :param Time:
        :param time_price:
        :return:
        """

        path_file = self.parent_dir+bill_name+".docx"

        if not self.check_dir(self.parent_dir):
            self.create_dir(self.parent_dir)

        bill_info = self.order_detail.select_detail_order(order_id)
        self.insert_row_and_value(self.origin_path, path_file, bill_info)

        infor  = [("","Thời gian", origin_time_price , Time, time_price),
                  ("","", "", 'Ngày', date),
                  ("","", "", 'Tổng', Total)
                  ]
        self.insert_row_and_value(path_file, path_file, infor)
        self.add_infor(path_file, path_file, "save", font_size=10, text=" " , option="end")
        self.show_bill_to_label(window, path_file)

        return path_file


    def manage_debtor(self, order_id,debtor, old_path, debtor_new_file):
        """
        if debtor exist => add debtor to the bill and save in the debtor directory and remove the initial bill file
        :param order_id:
        :param debtor:
        :param old_path:
        :param debtor_new_file:
        :return:
        """

        if not self.check_dir(self.deb_parent_dir):
            self.create_dir(self.deb_parent_dir)

        new_path = self.deb_parent_dir+debtor_new_file

        if len(debtor) > 1:

            infor = [(" ", " ", " ", "người nợ", debtor)]
            self.insert_row_and_value(old_path, new_path, infor)
            self.order_control.update_debtor(debtor, order_id)

            if os.path.exists(old_path):
                os.remove(old_path)
            else:
                print("The file does not exist")
            return  new_path





